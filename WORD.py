from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 9) # texto, ancho de letra

# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('negrita').bold = True # bold
# p.add_run(' and some ') # arial
# p.add_run('inclinada.').italic = True #italica

document.add_heading('Heading, level 1', level=1) #level = ancho de letra
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

#document.add_picture('utiles.png', width=Inches(5), height = Inches(2)) #path imagen, ancho de imagen, alto de imagen

records = (
    (3, '101', 'Spam', 'hola'),
    (7, '422', 'Eggs', 'adios'),
    (4, '631', 'Spam, spam, eggs, and spam', 'jodete')
)

table = document.add_table(rows=3, cols=4) #tabla filas, columnas
hdr_cells = table.rows[0].cells #craer primera fila
hdr_cells[0].text = 'Qty' #marcar columna de la fila antes creada
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
hdr_cells[3].text = 'saludo'
for qty, id, desc, saludo in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
    row_cells[3].text = saludo

document.add_page_break() #cerrar archivo

document.save('demo.docx')

#powerpoint
# from pptx import Presentation

# prs = Presentation()
# title_slide_layout = prs.slide_layouts[0]
# slide = prs.slides.add_slide(title_slide_layout)
# title = slide.shapes.title
# subtitle = slide.placeholders[1]

# title.text = "Hello, World!"
# subtitle.text = "python-pptx was here!"

# prs.save('test.pptx')