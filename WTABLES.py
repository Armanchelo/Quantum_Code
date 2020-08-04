import docx
import os
import numpy

doc = docx.Document()
titulo = input('Título: ')
doc.add_heading(titulo, 0)

def crearTabla():
    Tabla = []
    columnas = int(input('Número de columnas: '))
    filas = int(input('Número de filas: '))

    for f in range(filas):
        row = []
        for c in range(columnas):
            row.append(input(f'Columna {c + 1}, Fila {f + 1}: '))
        Tabla.append(row)

    tabla = doc.add_table(rows = filas, cols = len(Tabla[0]))
    tabla.style = 'Table Grid'

    for r in range(filas):
        fila = tabla.add_row().cells
        for c in range(columnas):
            fila[c].text = Tabla[r][c]

nombre = 'tabla.docx'
doc.save(nombre)
os.system(nombre)