import speech_recognition as sr
import wikipedia
from googletrans import Translator
import pyttsx3
from docx import Document
from math import pi, cos, sin
import win32api, win32con, win32gui
from chempy import balance_stoichiometry
from bs4 import BeautifulSoup
import requests
import wget
import datetime
import os
from os import listdir
from os.path import isfile, isdir
import wolframalpha
import webbrowser
try: import httplib
except: import http.client as httplib
import random

traductor = Translator()
r = sr.Recognizer()
voz = pyttsx3.init()
voz.setProperty('rate', 185)
document = Document()
word_folder = 'C:/Users/Usuario/Desktop/word/'
ironmanOriginal = r'C:\Users\Usuario\favorite ironman.jpg'
ironmanNew = r'C:\Users\Usuario\Iron-Man-Wallpaper.jpg'
creadores = [
    'Ángel Armando Chaparro Rocha', 'Javier Alejandro Vázquez Saenz', 
    'Ío Fernanda Resendez','Jorge Demian Muñoz Marín', 'en colaboración del doctor Adrián Domínguez Rodríguez'
]
wikipedia.set_lang('es')

buscadores = {
    'google': 'https://www.google.com/search?q=', 'youtube': 'https://www.youtube.com/results?search_query=', 
    'mercado libre': 'https://listado.mercadolibre.com.mx/', 'wikipedia': 'https://es.wikipedia.org/wiki/', 
    'google academic': 'https://scholar.google.com.mx/scholar?hl=es&as_sdt=0%2C5&q=', 'amazon': 'https://www.amazon.com.mx/s?k='
}

despedidas = ['chido', 'hasta luego', 'nos vemos', 'pórtate bien', 'suerte', 'bai', 'yo también me voy']

def hayInternet(timeout = 5):
    conn = httplib.HTTPConnection("www.google.com", timeout = timeout)
    try:
        conn.request("HEAD", "/")
        conn.close()
        return True
    except:
        conn.close()
        return False

def decir(texto, mostrar = True):
    reemplazar = {'PYTHON' : 'PAITON', 'MAFER' : 'MAFERR'}
    if mostrar == True:
        print(texto)
    for palabras in reemplazar:
        texto = texto.replace(palabras, reemplazar[palabras])
    voz.say(texto)
    voz.runAndWait()

def buscar(busqueda):
    Internet = hayInternet()
    if Internet:
        navegador = ''
        for buscador in buscadores:
            if busqueda.endswith('en ' + buscador): 
                busqueda = busqueda.replace('en ' + buscador, '').title()
                navegador = buscador
        if navegador != '':
            webbrowser.open(buscadores[navegador] + busqueda, new = 1, autoraise = True)
        else:
            decir('no sé usar ese buscador')
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

def despedida():
    decir(random.choices(despedidas)[0])
    exit()

def traducir(traduccion, language = 'es'):
    Internet = hayInternet()
    if Internet:
        if traduccion == 'texto':
            decir('introduce el texto')
            traduccion = input()
            traduccion = traductor.translate(traduccion, dest = language).text
            decir(traduccion)
        elif traduccion == 'documento':
            decir('introduce la ruta del archivo')
            ruta = input()
            traduccion = open(ruta, 'r').read()
            traduccion = traductor.translate(traduccion, dest = language).text
            decir(traduccion)
        else:
            return traductor.translate(traduccion, dest = language).text
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

appId = 'K86K7L-UJ8X6RKYVJ'
client = wolframalpha.Client(appId)
def wolfram(text=''):
    Internet = hayInternet()
    if Internet:
        res = client.query(text)
        if res['@success'] == 'false':
            decir('no tengo la respuesta')
        else:
            result = ''
            pod0 = res['pod'][0]
            pod1 = res['pod'][1]
            if (('definition' in pod1['@title'].lower()) or ('result' in  pod1['@title'].lower()) or (pod1.get('@primary','false') == 'true')):
                result = resolveListOrDict(pod1['subpod'])
                question = resolveListOrDict(pod0['subpod'])
                question = question.split('(')[0]
                answer = traducir(result)
                decir(answer)
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

def resolveListOrDict(variable):
  if isinstance(variable, list):
    return variable[0]['plaintext']
  else:
    return variable['plaintext']

class Wikipedia:
    def __init__(self, query, sentences = 0):
        self.query = query
        self.url = wikipedia.page(query).url
        self.url = query.title() + '.' + f'({datetime.datetime.now().year})' + '.' + self.url
        self.title = wikipedia.page(query).title
        self.result = wikipedia.summary(query)
        for n in range(1, 30):
            self.result = self.result.replace(f'[{n}]', '')
        for n in range(1000, 2100):
            self.result = self.result.replace(f'({n})', str(n))
        if sentences > 0:
            self.result = self.result.split('.')
            self.result = self.result[0:sentences]

def reiniciar(texto):
  if ' seconds' in texto or ' segundos' in texto:
    texto = texto.replace(' seconds', '')
    texto = texto.replace(' segundos', '')
    multiplicador = 1
  elif ' minutes' in texto or ' minutos' in texto:
    texto = texto.replace(' minutes', '')
    texto = texto.replace(' minutos', '')
    multiplicador = 60
  elif ' hours' in texto or ' horas' in texto:
    texto = texto.replace(' hours', '')
    texto = texto.replace(' horas', '')
    multiplicador = 3600
  time = int(texto) * multiplicador
  os.system(f'shutdown /r /t {time}')

def cancelarApagado():
  os.system('shutdown /a')

def apagar(texto):
  if ' seconds' in texto or ' segundos' in texto:
    texto = texto.replace(' seconds', '')
    texto = texto.replace(' segundos', '')
    multiplicador = 1
  elif ' minutes' in texto or ' minutos' in texto:
    texto = texto.replace(' minutes', '')
    texto = texto.replace(' minutos', '')
    multiplicador = 60
  elif ' hours' in texto or ' horas' in texto:
    texto = texto.replace(' hours', '')
    texto = texto.replace(' horas', '')
    multiplicador = 3600
  time = int(texto) * multiplicador
  os.system("shutdown /s /t " + str(time))

def cuentameSobre(busqueda):
    Internet = hayInternet()
    if Internet:
        texto = ''
        busqueda = Wikipedia(busqueda, 3)
        for result in busqueda.result:
            texto = texto + result + '.'
        texto = texto.replace('jul.', '')
        texto = texto.replace('greg.', '')
        decir(texto)
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

def crearReporte(busqueda):
    Internet = hayInternet()
    if Internet:
        busqueda = Wikipedia(busqueda)
        file_name = busqueda.query.title().replace(' ', '')
        document.add_paragraph(busqueda.url)
        document.add_paragraph(busqueda.result)
        document.add_page_break()
        document.save(word_folder + file_name + '.docx')
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

def limpiar_acentos(text):
	acentos = {'á': 'a', 'é': 'e', 'í': 'i', 'ó': 'o', 'ú': 'u', 'Á': 'A', 'É': 'E', 'Í': 'I', 'Ó': 'O', 'Ú': 'U'}
	for acen in acentos:
		if acen in text:
			text = text.replace(acen, acentos[acen])
	return text

def escuchar(lang, limite = 3):
    Internet = hayInternet()
    if Internet:
        decir('te escucho')
        with sr.Microphone() as source:
            if limite > 0:
                audio = r.listen(source, phrase_time_limit = limite)
            else:
                audio = r.listen(source)
            text = r.recognize_google(audio, language = lang)
            try:
                text = limpiar_acentos(text)
                return text.lower()
            except:
                pass
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

def Archivos(directorio):
    return [obj for obj in listdir(directorio) if isfile(directorio + obj)]

def tomarNota(numero):
    Internet = hayInternet()
    if Internet:
        archivos = Archivos('/Users/Usuario/LIZY/Notas/')
        NotaDisponible = False
        while NotaDisponible == False:
            Nota = 'nota ' + numero + '.txt'
            if Nota in archivos:
                decir('esa nota ya existe, pruebe otro número, o ¿desea reescribirla?')
                respuesta = escuchar('es', 0)
                if respuesta != 'reescribir': numero = respuesta 
            else:
                NotaDisponible = True
        decir('tomando nota')
        nota = escuchar('es', 0)
        try:
            nota = nota.replace(' punto final', '. ')
            nota = nota.replace(' punto y seguido', '.')
            nota = nota.replace(' coma', ',')
            nota = nota.replace(' punto y coma', ';')
            nota = nota.replace(' 2 puntos', ':')
            print(nota)
            newNote = open('C:\\Users\\Usuario\\LIZY\\Notas\\' + 'nota ' + numero + '.txt', 'x')
            newNote.write(nota)
        except:
            pass
    else:
        decir('no tengo internet, conéctate y prueba de nuevo')

def abrirNota(numero):
    Nota = 'nota ' + numero + '.txt'
    archivos = Archivos('/Users/Usuario/LIZY/Notas/')
    if Nota in archivos:
        texto = open('nota ' + numero + '.txt', 'r')
        nota = texto.read()
        decir(nota)
    else:
        decir('no has creado esa nota')

def esFuncion(funciones, comando):
    for funcion in funciones:
        if comando.startswith(funcion):
            return funcion
    return 'no es funcion'

def presentacion():
    hora = datetime.datetime.now().hour
    if hora >= 19 or hora <= 5: hora = 'hola, buenas noches'
    elif hora >= 6 or hora < 12: hora = 'hola, buenos días'
    else: hora = 'hola, buenas tardes'
    intro = hora + ' ,soy un proyecto elaborado en el colegio de bachilleres del estado de chihuahua plantel número 3, mi nombre es MAFER, y mis creadores son:'
    decir(intro)
    for creador in creadores:
        decir(creador)
    objetivo = 'Mi objetivo general es el de encontrar una nueva forma de impartir la enseñanza y el aprendizaje de una manera interactiva ya sea en un plantel o teleducación. Así poder incluir una actualización a la idea del estudio y recolección de datos dentro de estos. En particular me interesa ser un “asistente virtual” que interaccione con los alumnos y el laboratorio del plantel y facilitar la realización de prácticas. Reducir el tiempo en completar las tareas pendientes con ayuda de las TICS y PYTHON. Por último ser una aportación a la comunidad de PYTHON como un módulo de asistencia virtual'
    decir(objetivo)
    componentes = 'para estar aquí necesito de mi procesador, conexión a la nube, y opcionalmente arduino'
    decir(componentes)

def graficar():
    archivos = Archivos('/Users/Usuario/LIZY/')
    if 'LIZY_PLOT.py' in archivos:
        os.system('del LIZY_PLOT.py')
    fx = input('f(x) = ')
    if 'cos' in fx: fx = fx.replace('cos', 'np.cos')
    if 'sin' in fx: fx = fx.replace('sin', 'np.sin')
    if 'tan' in fx: fx = fx.replace('tan', 'np.tan')

    lima = input('límite inferior: ')
    limb = input('límite superior: ')

    with open('LIZY_PLOT.py', 'w') as fl:
        print('from matplotlib import pyplot as plt', file = fl)
        print('import numpy as np', file = fl)
        print('from math import *', file = fl)
        print('def PLOT():', file = fl)
        print(f'    x = np.arange({lima}, {limb}, 0.01)', file = fl)
        print(f'    y = {fx}', file = fl)
        print('    plt.plot(x, y)', file = fl)
        print('    plt.xlabel("X")', file = fl)
        print('    plt.ylabel("Y")', file = fl)
        print(f'    plt.title("{fx}")', file = fl)
        print('    plt.grid()', file = fl)
        print('    plt.show()', file = fl)
    
    import LIZY_PLOT as lp
    lp.PLOT()

def setWallpaper(path):
    key = win32api.RegOpenKeyEx(win32con.HKEY_CURRENT_USER,"Control Panel\\Desktop",0,win32con.KEY_SET_VALUE)
    win32gui.SystemParametersInfo(win32con.SPI_SETDESKWALLPAPER, path, 1+2)

def chemBalance():
    formula = input('fórmula: ')
    formula = formula.upper()
    formula = formula.replace(' ', '')
    formula = formula.split('=')
    reactivos = formula[0].split('+')
    productos = formula[1].split('+')
    reac, prod = balance_stoichiometry(reactivos, productos)
    reacs = ''
    index = 0
    for reactivo in reactivos:
        if reac.get(reactivo) != 1:
            reacs += str(reac.get(reactivo))
        reacs += reactivo
        if index < len(reactivos) - 1:
            reacs += ' + '
        index += 1
    reacs += ' = '
    index = 0
    for producto in productos:
        if prod.get(producto) != 1:
            reacs += str(prod.get(producto))
        reacs += producto
        if index < len(productos) - 1:
            reacs += ' + '
        index += 1
    print(reacs)